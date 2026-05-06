#!/usr/bin/env python3
"""
Build a complete Word .docx from the markdown manuscript draft + appendix +
references, embedding tables (Tables 4-7 inline plus all appendix tables) and
inserting figure references.
"""
import re
from pathlib import Path
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/gujaeseo/Documents/projects/yeonseo/ijim_research_2026")
MR   = ROOT / "manuscript_revised"
RESP = ROOT / "response"
FIG  = ROOT / "output" / "figures"
TBL  = ROOT / "output" / "tables"

doc = docx.Document()
n_style = doc.styles['Normal']
n_style.font.name = 'Times New Roman'
n_style.font.size = Pt(11)

# Helper: parse markdown tables (rough but works for our format)
def md_to_table(doc, md_text):
    """Parse markdown table into a docx table."""
    lines = [l.strip() for l in md_text.strip().split("\n") if l.strip()]
    rows = [r for r in lines if r.startswith("|")]
    if len(rows) < 2: return False
    sep_idx = next((i for i, r in enumerate(rows) if re.match(r"^\|[\s\-:|]+\|$", r)), -1)
    if sep_idx == -1: return False
    headers = [c.strip() for c in rows[0].strip("|").split("|")]
    body_rows = [[c.strip() for c in r.strip("|").split("|")] for r in rows[sep_idx+1:]]
    n_cols = len(headers)
    t = doc.add_table(rows=1+len(body_rows), cols=n_cols)
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(body_rows):
        for ci, cell in enumerate(row[:n_cols]):
            t.rows[ri+1].cells[ci].text = cell
    return True

def add_inline_md(doc, txt):
    """Add a paragraph treating a few inline markers (bold, italic, code)."""
    if not txt:
        doc.add_paragraph(""); return
    p = doc.add_paragraph()
    pos = 0
    parts = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)", txt)
    for pt in parts:
        if not pt: continue
        if pt.startswith("***") and pt.endswith("***"):
            r = p.add_run(pt[3:-3]); r.bold = True; r.italic = True
        elif pt.startswith("**") and pt.endswith("**"):
            r = p.add_run(pt[2:-2]); r.bold = True
        elif pt.startswith("*") and pt.endswith("*") and len(pt) > 2:
            r = p.add_run(pt[1:-1]); r.italic = True
        elif pt.startswith("`") and pt.endswith("`"):
            r = p.add_run(pt[1:-1]); r.font.name = "Courier New"
        else:
            p.add_run(pt)

def render_md(doc, md_path):
    """Render a markdown file into the doc, with table block detection."""
    text = Path(md_path).read_text()
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Heading
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        # Table block (gather contiguous '|' lines)
        elif line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            if not md_to_table(doc, "\n".join(tbl_lines)):
                for tl in tbl_lines: doc.add_paragraph(tl)
            continue
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
        elif line.strip() == "---":
            doc.add_paragraph("─" * 80)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            add_inline_md(doc, line)
        i += 1

# ----------- Build -----------
doc.add_heading("Does AI Benefit All Platform Stakeholders? Evidence from Algorithmic Task Assignment in Food Delivery", 0)
p = doc.add_paragraph("Revised manuscript — generated 2026-05-06"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sections 1-7
for n in range(1, 8):
    fs = list(MR.glob(f"section_{n}_*.md"))
    if not fs: continue
    doc.add_page_break()
    render_md(doc, fs[0])

# Appendix
doc.add_page_break()
render_md(doc, MR / "appendix.md")

# Insert figure images at end of appendix
doc.add_page_break()
doc.add_heading("Figures (with captions)", level=1)
fig_files = [
    ("Figure A1. Event-study coefficients on productivity (DID, week-level leads/lags).",
     "05_event_study_productivity.png"),
    ("Figure A2. Dose-response: AI assignment intensity vs. productivity (post-period treated rider-days).",
     "07_dose_response.png"),
    ("Figure A3. Pre-vs-post productivity density (matched 336 riders).",
     "09_inequality_density.png"),
    ("Figure A4. Learning dynamics by proficiency group (week-level treatment coefficients).",
     "10_learning_curves_by_group.png"),
    ("Figure A5. Worker-customer linkage: stack productivity vs customer waiting time.",
     "11_worker_customer_linkage.png"),
    ("Figure A6. Sample representativeness: matched analytic sample vs broader preexist Busan riders.",
     "12_sample_representativeness.png"),
    ("Figure A7. Long-term productivity event-study (monthly coefficients, Sep 2020 – Jan 2021).",
     "13_long_term_event_study.png"),
]
for caption, fname in fig_files:
    fp = FIG / fname
    if fp.exists():
        try:
            doc.add_picture(str(fp), width=Inches(6.0))
            cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs: run.italic = True
        except Exception as e:
            doc.add_paragraph(f"[Figure not embeddable: {fname} — {e}]")

# References
doc.add_page_break()
render_md(doc, MR / "references.md")

out = MR / "manuscript_revised_v2_full.docx"
doc.save(out)
print(f"Saved: {out}  size={out.stat().st_size} bytes")

# Response letter docx (regen)
rdoc = docx.Document()
rdoc.styles['Normal'].font.name = 'Times New Roman'
rdoc.styles['Normal'].font.size = Pt(11)
render_md(rdoc, RESP / "response_letter_v1.md")
rout = RESP / "response_letter_v1.docx"
rdoc.save(rout)
print(f"Saved: {rout}  size={rout.stat().st_size} bytes")
